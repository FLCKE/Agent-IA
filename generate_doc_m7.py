import os
import subprocess
import re
from docx import Document
from docx.shared import Pt, Inches

def sanitize_text(text):
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@~])')
    text = ansi_escape.sub('', text)
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
    return text

def run_lab(lab_file):
    try:
        result = subprocess.run(
            ["python", lab_file],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300
        )
        output = result.stdout
        if result.stderr:
            output += "\n[ERREUR]:\n" + result.stderr
        return sanitize_text(output.strip())
    except subprocess.TimeoutExpired:
        return "[TIMEOUT] L'exécution a pris trop de temps."
    except Exception as e:
        return f"[ERREUR] {str(e)}"

def add_code_block(doc, code):
    p = doc.add_paragraph(code)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)

def main():
    doc = Document()
    doc.add_heading("Module 7 - Les Agents Autonomes", 0)
    
    p = doc.add_paragraph()
    p.add_run("Étudiant: FRANCISCO Louis-Carlos\n").bold = True
    p.add_run("Livrables des Labs 1 à 6")

    labs = [
        {"name": "Labo 1 — Agent orienté objectif", "file": "Module7_FRANCISCO_LOUISCARLOS/lab1.py"},
        {"name": "Labo 2 — Boucle d'autonomie", "file": "Module7_FRANCISCO_LOUISCARLOS/lab2.py"},
        {"name": "Labo 3 — Auto-réflexion et auto-correction", "file": "Module7_FRANCISCO_LOUISCARLOS/lab3.py"},
        {"name": "Labo 4 — Sécurité et garde-fous", "file": "Module7_FRANCISCO_LOUISCARLOS/lab4.py"},
        {"name": "Labo 5 — Évaluation de l'autonomie", "file": "Module7_FRANCISCO_LOUISCARLOS/lab5.py"},
        {"name": "Labo 6 — Auto-vérification (SelfCheck)", "file": "Module7_FRANCISCO_LOUISCARLOS/lab6.py"},
    ]

    for lab in labs:
        doc.add_page_break()
        doc.add_heading(lab["name"], level=1)
        
        filepath = lab["file"]
        doc.add_heading("Code source", level=2)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                code_content = f.read()
            add_code_block(doc, code_content)
        except Exception as e:
            doc.add_paragraph(f"Erreur de lecture: {e}")

        doc.add_heading("Résultat (Sortie console)", level=2)
        print(f"Exécution de {filepath}...")
        output = run_lab(filepath)
        add_code_block(doc, output)

    doc.add_page_break()
    doc.add_heading("Synthèse et Analyse", level=1)
    doc.add_paragraph(
        "Ce module a permis d'explorer la transition vers une autonomie réelle des agents IA. "
        "L'utilisation de boucles itératives (Plan-Act-Reflect) montre une nette amélioration de la qualité des réponses par rapport à un simple appel unique. "
        "Les mécanismes de sécurité implémentés dans le Lab 4 sont cruciaux pour éviter les dérives et les coûts imprévus. "
        "L'évaluation via KPI (Lab 5) et la double vérification (Lab 6) assurent un contrôle rigoureux sur les systèmes autonomes."
    )

    output_path = "Module7_FRANCISCO_LOUISCARLOS/Module7_FRANCISCO_LouisCarlos.docx"
    doc.save(output_path)
    print(f"Document généré: {output_path}")

if __name__ == "__main__":
    main()
