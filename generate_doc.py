import os
import subprocess
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def sanitize_text(text):
    # Remove ANSI escape sequences
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    text = ansi_escape.sub('', text)
    # Remove control characters except newline and tab
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
    return text

def run_lab(lab_file):
    try:
        result = subprocess.run(
            ["python", lab_file],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300
        )
        output = result.stdout
        if result.stderr:
            output += "\n[ERREUR]:\n" + result.stderr
        return sanitize_text(output.strip())
    except subprocess.TimeoutExpired:
        return "[TIMEOUT] L'exécution a pris trop de temps."
    except Exception as e:
        return f"[ERREUR] {str(e)}"

def add_code_block(doc, code):
    p = doc.add_paragraph(code)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)

def main():
    doc = Document()
    doc.add_heading("Module 6 - Les Systèmes Multi-Agents", 0)
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run("Étudiant: FRANCISCO Louis-Carlos\n").bold = True
    p.add_run("Livrables des Labs 1 à 5")

    labs = [
        {"name": "Labo 1 — Spécialisation d'agents", "file": "Module6_FRANCISCO_LOUISCARLOS/lab1.py"},
        {"name": "Labo 2 — Architecture Manager-Worker", "file": "Module6_FRANCISCO_LOUISCARLOS/lab2.py"},
        {"name": "Labo 3 — Communication et mémoire partagée", "file": "Module6_FRANCISCO_LOUISCARLOS/lab3.py"},
        {"name": "Labo 4 — Mini CrewAI : équipe collaborative", "file": "Module6_FRANCISCO_LOUISCARLOS/lab4.py"},
        {"name": "Labo 5 — Consensus et résolution de conflits", "file": "Module6_FRANCISCO_LOUISCARLOS/lab5.py"},
    ]

    for lab in labs:
        doc.add_page_break()
        doc.add_heading(lab["name"], level=1)
        
        filepath = lab["file"]
        
        # Read source code
        doc.add_heading("Code source", level=2)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                code_content = f.read()
            add_code_block(doc, code_content)
        except Exception as e:
            doc.add_paragraph(f"Erreur de lecture du code: {e}")

        # Run and capture output
        doc.add_heading("Résultat (Sortie console)", level=2)
        print(f"Exécution de {filepath}...")
        output = run_lab(filepath)
        add_code_block(doc, output)
        print(f"Terminé pour {filepath}.")

    # Add Mini-analyse for Lab 5 as requested by the assignment
    doc.add_heading("Mini-analyse des méthodes de consensus", level=2)
    analyse = (
        "Le vote majoritaire est la méthode la plus simple et souvent efficace pour des décisions binaires, "
        "mais il ne prend pas en compte la certitude ou l'expertise de chaque agent. "
        "Le score de confiance permet de pondérer les réponses, privilégiant l'agent le plus sûr de lui, "
        "ce qui est utile si les agents savent évaluer leur propre fiabilité. "
        "Enfin, l'agent arbitre permet une réflexion plus nuancée, capable de juger non seulement la réponse "
        "mais aussi la justification apportée par chaque agent, ce qui est idéal pour des questions complexes "
        "comme la conscience de l'IA abordée dans ce labo."
    )
    doc.add_paragraph(analyse)

    output_path = "Module6_FRANCISCO_LOUISCARLOS/Module6_FRANCISCO_LouisCarlos.docx"
    doc.save(output_path)
    print(f"Document généré avec succès: {output_path}")

if __name__ == "__main__":
    main()
